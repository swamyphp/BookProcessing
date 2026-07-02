#!/usr/bin/env python3
import sys, os, json
from docx import Document

def find_manuscript(folder):
    for root, dirs, files in os.walk(folder):
        for f in files:
            if f.lower()=='manuscript.docx':
                return os.path.join(root,f)
    return None

def extract_chapters(doc_path):
    doc = Document(doc_path)
    chapters = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text: continue
        # simple heuristic: lines starting with 'Chapter' or styled as Heading 1
        if text.lower().startswith('chapter'):
            chapters.append(text)
    # fallback: count headings
    if not chapters:
        for p in doc.paragraphs:
            if p.style.name.lower().startswith('heading'):
                chapters.append(p.text.strip())
    return chapters

def main():
    if len(sys.argv)<3:
        print(json.dumps({'status':'error','error':'args'}))
        return
    folder = sys.argv[1]
    short = sys.argv[2]
    m = find_manuscript(folder)
    if not m:
        print(json.dumps({'status':'error','error':'manuscript missing'}))
        return
    chapters = extract_chapters(m)
    out = {'status':'success','book':short,'book_title':short,'chapters':len(chapters),'chapter_titles':chapters}
    print(json.dumps(out))

if __name__=='__main__':
    main()
